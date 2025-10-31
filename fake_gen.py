import os
import random
import string
import requests
import zipfile
from faker import Faker
from docx import Document
from reportlab.pdfgen import canvas

# ---------------- Configuration ----------------
# num_docs = 10        # number of documents
# num_images = 10      # number of images
# num_downloads = 5
image_width = 800
image_height = 600


types = ['txt', 'pdf', 'docx']  # document types
faker = Faker()

# ---------------- Italian word list ----------------
words = [
    "casa", "sole", "luna", "mare", "montagna", "fiume", "bosco", "cielo",
    "stella", "nuvola", "fiore", "strada", "albero", "vento", "pioggia",
    "neve", "fuoco", "terra", "sogno", "ombra", "luce", "giardino",
    "apple", "orange", "banana", "grape", "river", "mountain", "forest", 
    "sky", "cloud", "ocean", "desert", "star", "moon", "sun", "flower",
    "tree", "stone", "road", "city", "village", "house", "garden", "wind",
    "rain", "snow", "fire", "earth", "dream", "shadow", "light"
]

sites = [
    'unipd.it',
    'google.com',
    'apple.com',
    'facebook.com',
    'x.com',
    'instagram.com',
    'wiki.archlinux.org',
    'raceup.it',
    'science4all.it'
]

# ---------------- Helper Functions ----------------
def random_name():
    word1 = random.choice(words)
    word2 = random.choice(words)
    number = random.randint(10, 99)
    return f"{word1}_{word2}_{number}"

def create_txt_file(path):
    paragraphs = [faker.paragraph(nb_sentences=random.randint(3,7)) for _ in range(random.randint(1,3))]
    with open(path, "w") as f:
        f.write("\n\n".join(paragraphs))

def create_pdf_file(path):
    c = canvas.Canvas(path)
    text = c.beginText(50, 800)
    for _ in range(random.randint(5,15)):
        text.textLine(faker.sentence())
    c.drawText(text)
    c.save()

def create_docx_file(path, random_name='Fake Document'):
    doc = Document()
    doc.add_heading(random_name, level=0)
    for _ in range(random.randint(3,7)):
        doc.add_paragraph(faker.paragraph())
    doc.save(path)

def download_image(path):
    url = f"https://picsum.photos/{image_width}/{image_height}"
    try:
        img_data = requests.get(url).content
        with open(path, "wb") as f:
            f.write(img_data)
    except Exception as e:
        print(f"⚠️ Failed to download {path}: {e}")

def download_page(path, url):
    try:
        html = requests.get('https://' + url).content
        with open(path, "wb") as f:
            f.write(html)
    except Exception as e:
        print(f"⚠️ Failed to download {path}: {e}")

# ---------------- Main ----------------
def populate_files(root, num_images=10, num_docs=10, num_downloads=5):
    output_root = root
    documents_dir = os.path.join(output_root, "Documents")
    pictures_dir = os.path.join(output_root, "Pictures")
    downloads_dir = os.path.join(output_root, "Downloads")

    os.makedirs(documents_dir, exist_ok=True)
    os.makedirs(pictures_dir, exist_ok=True)
    os.makedirs(downloads_dir, exist_ok=True)

    # Create documents
    print(f"Creating {num_docs} random documents...")
    for _ in range(num_docs):
        ext = random.choice(types)
        name = random_name()
        filename = f"{name}.{ext}"
        full_path = os.path.join(documents_dir, filename)
        
        if ext == 'txt':
            create_txt_file(full_path)
        elif ext == 'pdf':
            create_pdf_file(full_path)
        elif ext == 'docx':
            create_docx_file(full_path, name)

        print(f"  → {filename}")

    # Download images
    print(f"\nDownloading {num_images} random images...")
    for _ in range(num_images):
        filename = f"{random_name()}.jpg"
        full_path = os.path.join(pictures_dir, filename)
        download_image(full_path)
        print(f"  → {filename}")

    # Download pages
    print(f"Downloading {num_downloads} random pages...")
    random_sites = random.sample(sites, num_downloads)
    for site in random_sites:
        filename = f"{random_name()}_{site}.html"
        full_path = os.path.join(downloads_dir, filename)
        download_page(full_path, site)
        print(f"  → {filename}")

    print(f"\n✅ Populated '{output_root}' with documents and images.")
