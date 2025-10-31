import random
import io
import requests
import tempfile
import shutil
import zipfile
import os
from faker import Faker
from docx import Document
from docx.shared import Inches
from PIL import Image

def download_random_image(width=400, height=300):
    """Download a random image from picsum.photos."""
    url = f"https://picsum.photos/{width}/{height}"
    response = requests.get(url)
    if response.status_code == 200:
        return io.BytesIO(response.content)
    return None

def generate_random_doc(filename="random_document.docx", num_paragraphs=5, num_images=4, corrupt_document_xml=True, corrupt_content_types=False):
    fake = Faker()
    doc = Document()

    # Random title
    title = fake.sentence(nb_words=random.randint(3, 6)).title()
    doc.add_heading(title, level=0)

    # Generate paragraphs and random images in between
    for i in range(num_paragraphs):
        paragraph = fake.paragraph(nb_sentences=random.randint(3, 8))
        doc.add_paragraph(paragraph)

        # Occasionally insert an image
        if random.random() < 0.5 and num_images > 0:
            img_data = download_random_image()
            if img_data:
                doc.add_picture(img_data, width=Inches(3.5))
                num_images -= 1

    # Add a closing paragraph
    doc.add_paragraph(fake.paragraph(nb_sentences=5))

    temp_original = tempfile.mktemp(prefix="docx_original_")
    docx_original_path = os.path.join(temp_original, filename)
    doc.save(docx_original_path)

    """
    Create a copy of src_docx named out_docx with certain XML parts damaged.
    We always preserve 'word/media/*' so images remain retrievable.
    """
    # Make temp dir to extract zip
    tmpdir = tempfile.mkdtemp(prefix="docx_corrupt_")
    try:
        with zipfile.ZipFile(docx_original_path, 'r') as zin:
            zin.extractall(tmpdir)

        # Paths to potentially corrupt
        doc_xml_path = os.path.join(tmpdir, "word", "document.xml")
        content_types_path = os.path.join(tmpdir, "[Content_Types].xml")

        # 1) Corrupt document.xml by truncating it and injecting invalid bytes (if available)
        if corrupt_document_xml and os.path.exists(doc_xml_path):
            print("[*] Corrupting word/document.xml ...")
            with open(doc_xml_path, "rb") as f:
                original = f.read()

            # Strategy: keep the XML prolog and first ~200 bytes, then insert garbage and cut off
            prolog_end = original.find(b'>')  # naive end of <?xml ... ?> maybe
            if prolog_end == -1:
                prolog_end = 0
            keep = original[:max(prolog_end + 1, 200)]
            garbage = b"\n<!-- CORRUPTED -->\n" + b"\x00\xff\x00\xff" * 10 + b"UNREADABLE-BY-XML"
            truncated = keep + garbage  # no closing tags -> invalid XML
            # Write back
            with open(doc_xml_path, "wb") as f:
                f.write(truncated)
            print("[+] word/document.xml replaced with truncated/garbage content.")

        # 2) Optionally corrupt [Content_Types].xml by breaking a tag name (Word will often fail to load)
        if corrupt_content_types and os.path.exists(content_types_path):
            print("[*] Corrupting [Content_Types].xml ...")
            with open(content_types_path, "rb") as f:
                ct = f.read()
            # Replace a common tag name 'Types' with 'Typ<>es' to break parsing
            broken = ct.replace(b"Types", b"Typ<>es", 1)
            # If the replace didn't change anything, just truncate
            if broken == ct:
                broken = ct[:120] + b"<!-- truncated -->"
            with open(content_types_path, "wb") as f:
                f.write(broken)
            print("[+] [Content_Types].xml corrupted.")

        # 3) Repack the folder into a new zip (docx)
        # Use ZIP_DEFLATED for compatibility
        with zipfile.ZipFile(filename, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
            # Walk original extracted structure and write files preserving relative paths
            for root, dirs, files in os.walk(tmpdir):
                for file in files:
                    fullpath = os.path.join(root, file)
                    relpath = os.path.relpath(fullpath, tmpdir)
                    # Write file into zip with the same relative path
                    zout.write(fullpath, relpath)
        print(f"[+] Wrote corrupted DOCX: {filename}")

        # Extra note: ensure media files are still present
        with zipfile.ZipFile(filename, 'r') as z:
            media_files = [n for n in z.namelist() if n.startswith("word/media/")]
            print(f"[i] word/media contains {len(media_files)} file(s):")
            for m in media_files:
                print("   -", m)

    finally:
        # Clean up temp dir
        shutil.rmtree(tmpdir)
        shutil.rmtree(temp_original)

def generated_corrupted(path):
    generate_random_doc(path)
    corrupt_docx_preserve_media()

if __name__ == "__main__":
    original = "dummy.docx"
    corrupted = "corrupt.docx"

    generate_random_doc(original,10,4)

    corrupt_docx_preserve_media(original, corrupted, corrupt_document_xml=True, corrupt_content_types=False)
